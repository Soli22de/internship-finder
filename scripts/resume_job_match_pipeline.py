import json
import os
import re
from collections import Counter
from datetime import datetime, timedelta
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document


RESUME_PATH = r"data\张靖恒-秋招大厂数据岗简历优化版.docx"
TARGET_PATH = r"outputs\reports\internship_target_jobs.csv"
MASTER_PATH = r"outputs\reports\internship_all_master.csv"
FALSE_NEGATIVE_PATH = r"outputs\reports\cohort27_false_negative_audit.csv"
REPORT_DIR = r"outputs\reports"
DATA_DIR = r"data"

TODAY = datetime.now().date()

SKILL_KEYWORDS = [
    "python",
    "sql",
    "pyspark",
    "spark",
    "hive",
    "airflow",
    "etl",
    "数据建模",
    "数据仓库",
    "数仓",
    "ab测试",
    "a/b",
    "可视化",
    "power bi",
    "tableau",
    "机器学习",
    "统计",
    "回归",
    "分类",
    "xgboost",
    "推荐",
    "策略",
    "增长",
    "埋点",
    "指标体系",
    "实验",
    "llm",
    "提示词",
    "nlp",
    "linux",
    "java",
    "go",
    "shell",
    "postgresql",
]

SCENE_KEYWORDS = ["电商", "本地生活", "增长", "策略", "推荐", "用户", "商业分析", "金融", "零售", "内容", "广告", "外卖", "酒旅", "物流"]
REMOTE_KEYWORDS = ["远程", "异地可投", "上海/北京", "全国", "多地", "base上海"]
FULLTIME_BLOCK_KEYWORDS = ["全职经验", "社招", "3年以上", "工作经验", "非实习"]
RETURN_OFFER_KEYWORDS = ["留用", "转正", "offer", "提前批", "校招"]
TERM_ALIASES = {
    "a/b": "ab测试",
    "ab test": "ab测试",
    "a-b test": "ab测试",
    "a/b test": "ab测试",
    "etl链路": "etl",
    "extract transform load": "etl",
    "数仓": "数据仓库",
    "data warehouse": "数据仓库",
    "bi": "可视化",
    "powerbi": "power bi",
    "tableau bi": "tableau",
    "nlp": "nlp",
    "large language model": "llm",
}
RESPONSIBILITY_HINTS = ["负责", "参与", "推动", "搭建", "维护", "优化", "协同", "设计", "落地", "执行"]
ABILITY_HINTS = ["熟悉", "掌握", "精通", "能力", "技能", "经验", "了解", "会", "能够", "具备"]
RESULT_HINTS = ["提升", "增长", "降低", "优化", "产出", "结果", "效果", "指标", "roi", "%", "效率"]


def ensure_dirs() -> None:
    os.makedirs(REPORT_DIR, exist_ok=True)
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(os.path.join(DATA_DIR, "张靖恒-秋招大厂数据岗简历_分岗位族定制版文件夹"), exist_ok=True)


def read_resume_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def extract_resume_profile(text: str) -> Dict[str, object]:
    low = text.lower()
    skills = [k for k in SKILL_KEYWORDS if k in low]
    scenes = [k for k in SCENE_KEYWORDS if k in text]
    quant_numbers = re.findall(r"\d+\.?\d*%?|\d+万|\d+千|\d+亿", text)
    intern_days = 4
    intern_months = 3
    grad_year = 2027
    m_day = re.search(r"每周.{0,4}?(\d)天", text)
    if m_day:
        intern_days = int(m_day.group(1))
    m_month = re.search(r"(\d+)\s*个?月", text)
    if m_month:
        intern_months = int(m_month.group(1))
    m_grad = re.search(r"(202[5-9])届", text)
    if m_grad:
        grad_year = int(m_grad.group(1))
    return {
        "resume_text": text,
        "skills": skills,
        "scenes": scenes,
        "quant_count": len(quant_numbers),
        "intern_days": intern_days,
        "intern_months": intern_months,
        "grad_year": grad_year,
    }


def normalize_publish_date(v: str) -> datetime:
    s = str(v).strip()
    if not s:
        return datetime(2000, 1, 1)
    if s.isdigit():
        iv = int(s)
        if iv > 10_000_000_000:
            return datetime.fromtimestamp(iv / 1000)
        if iv > 1_000_000_000:
            return datetime.fromtimestamp(iv)
    for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"]:
        try:
            return datetime.strptime(s[: len(fmt)], fmt)
        except Exception:
            pass
    return datetime(2000, 1, 1)


def parse_requirement_num(text: str, kind: str) -> int:
    t = str(text)
    if kind == "days":
        m = re.search(r"每周.{0,6}?(\d)\s*天", t)
        return int(m.group(1)) if m else 0
    m = re.search(r"(\d+)\s*个?月", t)
    return int(m.group(1)) if m else 0


def extract_deadline_date(v: str) -> datetime:
    s = str(v).strip()
    if not s:
        return datetime(2099, 12, 31)
    m = re.search(r"(20\d{2})[./-](\d{1,2})[./-](\d{1,2})", s)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception:
            return datetime(2099, 12, 31)
    return datetime(2099, 12, 31)


def role_family(name: str, jd: str) -> str:
    t = f"{name} {jd}".lower()
    if any(k in t for k in ["数据开发", "数仓", "etl", "hive", "spark", "数据工程", "大数据"]):
        return "数据开发/数据工程"
    if any(k in t for k in ["策略", "增长", "运营策略", "商业策略"]):
        return "策略分析/增长分析"
    if any(k in t for k in ["数据科学", "算法", "模型", "机器学习", "推荐"]):
        return "数据科学/算法分析"
    return "数据分析/商业分析"


def hard_screen_row(row: pd.Series, profile: Dict[str, object]) -> Tuple[bool, str]:
    reasons = []
    text = f"{row.get('name','')} {row.get('jd_raw','')} {row.get('requirement','')} {row.get('degree','')} {row.get('city','')}"
    constraints = extract_time_constraints(text, default_grad_year=int(profile["grad_year"]))
    if "硕士" in str(row.get("degree", "")) and "本科" not in str(row.get("degree", "")):
        reasons.append("学历仅限硕士及以上")
    if constraints["grad_year_required"] > int(profile["grad_year"]):
        reasons.append("毕业年份不匹配2027届")
    city = str(row.get("city", ""))
    if ("上海" not in city) and not any(k in text for k in REMOTE_KEYWORDS):
        reasons.append("地域不匹配上海/异地可投")
    day_req = int(constraints["intern_days_required"])
    if day_req and day_req > int(profile["intern_days"]):
        reasons.append("每周实习天数不满足")
    month_req = int(constraints["intern_months_required"])
    if month_req and month_req > int(profile["intern_months"]):
        reasons.append("实习时长不满足")
    if int(constraints["experience_years_required"]) >= 2:
        reasons.append("经验年限门槛偏高")
    if any(k in text for k in FULLTIME_BLOCK_KEYWORDS):
        reasons.append("要求全职经验")
    deadline = extract_deadline_date(row.get("deadline", ""))
    if deadline.date() < TODAY:
        reasons.append("已过截止日期")
    pass_flag = len(reasons) == 0
    return pass_flag, "；".join(reasons) if reasons else "通过硬门槛"


def split_sentences(text: str) -> List[str]:
    return [x.strip() for x in re.split(r"[。；\n]", str(text)) if x.strip()]


def normalize_term(term: str) -> str:
    t = str(term).strip().lower()
    if not t:
        return ""
    return TERM_ALIASES.get(t, t)


def normalized_skill_hits(text: str) -> List[str]:
    low = str(text).lower()
    normalized = set()
    for k in SKILL_KEYWORDS:
        if k in low:
            normalized.add(normalize_term(k))
    for alias, canonical in TERM_ALIASES.items():
        if alias in low:
            normalized.add(canonical)
    return sorted(x for x in normalized if x)


def extract_sentence_mapping(jd: str) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for s in split_sentences(jd):
        low = s.lower()
        role = "responsibility"
        if any(h in low for h in ABILITY_HINTS):
            role = "ability"
        if any(h in low for h in RESULT_HINTS):
            role = "result"
        if role == "responsibility" and any(h in low for h in RESPONSIBILITY_HINTS):
            role = "responsibility"
        rows.append(
            {
                "sentence": s,
                "role": role,
                "skills": "、".join(normalized_skill_hits(s)),
            }
        )
    return rows


def extract_time_constraints(text: str, default_grad_year: int = 2027) -> Dict[str, object]:
    t = str(text)
    days = 0
    months = 0
    exp_years = 0
    grad_year = default_grad_year
    arrival = ""
    m_day = re.search(r"每周.{0,6}?(\d)\s*天", t)
    if m_day:
        days = int(m_day.group(1))
    m_month = re.search(r"(\d+)\s*个?月", t)
    if m_month:
        months = int(m_month.group(1))
    m_exp = re.search(r"(\d+)\s*年.{0,4}经验", t)
    if m_exp:
        exp_years = int(m_exp.group(1))
    m_grad = re.search(r"(202[5-9])届", t)
    if m_grad:
        grad_year = int(m_grad.group(1))
    if "尽快到岗" in t or "立即到岗" in t:
        arrival = "immediate"
    elif "一周内到岗" in t:
        arrival = "within_1_week"
    elif "两周内到岗" in t or "2周内到岗" in t:
        arrival = "within_2_weeks"
    return {
        "intern_days_required": days,
        "intern_months_required": months,
        "experience_years_required": exp_years,
        "grad_year_required": grad_year,
        "arrival_required": arrival,
    }


def extract_must_plus_keywords(jd: str) -> Tuple[List[str], List[str]]:
    must = set()
    plus = set()
    low = str(jd).lower()
    sentences = split_sentences(low)
    for s in sentences:
        hit = normalized_skill_hits(s)
        if not hit:
            continue
        if any(x in s for x in ["必须", "必备", "要求", "掌握", "熟练", "精通"]):
            must.update(hit)
        elif any(x in s for x in ["优先", "加分", "更佳", "最好"]):
            plus.update(hit)
    if not must:
        top = normalized_skill_hits(low)[:3]
        must.update(top)
    return sorted(must), sorted(plus - must)


def match_score(row: pd.Series, profile: Dict[str, object]) -> Dict[str, object]:
    jd = f"{row.get('name','')} {row.get('jd_raw','')} {row.get('requirement','')} {row.get('raw_tags','')} {row.get('recruit_type','')}"
    jd_low = jd.lower()
    must, plus = extract_must_plus_keywords(jd)
    resume_skill = set(profile["skills"])
    must_hit = [k for k in must if k in resume_skill]
    plus_hit = [k for k in plus if k in resume_skill]
    must_rate = len(must_hit) / max(1, len(must))
    ats_score = 35 * must_rate
    scene_hit = [k for k in profile["scenes"] if k in jd]
    scene_score = 30 * (min(1.0, len(scene_hit) / 2) if scene_hit else 0)
    plus_rate = len(plus_hit) / max(1, len(plus)) if plus else 0.6
    evidence_base = min(1.0, float(profile["quant_count"]) / 12)
    evidence_score = 20 * (0.6 * evidence_base + 0.4 * plus_rate)
    pub = normalize_publish_date(row.get("publish_time", ""))
    days_old = max(0, (datetime.now() - pub).days)
    freshness = 1.0 if days_old <= 15 else (0.7 if days_old <= 30 else 0.3)
    return_signal = 1.0 if any(k in jd for k in RETURN_OFFER_KEYWORDS) else 0.5
    link_ok = 1.0 if str(row.get("link_status", "OK")) in {"OK", "UNKNOWN"} else 0.2
    value_score = 15 * (0.5 * freshness + 0.3 * return_signal + 0.2 * link_ok)
    total = round(ats_score + scene_score + evidence_score + value_score, 2)
    miss = [k for k in must if k not in resume_skill][:3]
    hit_points = "、".join((must_hit[:2] + scene_hit[:1])[:3]) or "基础关键词命中"
    gap_points = "、".join(miss) or "无明显必选关键词缺口"
    fam = role_family(str(row.get("name", "")), str(row.get("jd_raw", "")))
    if total >= 90 and must_rate >= 0.999 and any(k in jd for k in RETURN_OFFER_KEYWORDS):
        tier = "A1"
    elif total >= 80 and must_rate >= 0.999:
        tier = "A2"
    elif total >= 70 and len(miss) <= 1:
        tier = "B1"
    elif total >= 60 and len(miss) <= 2:
        tier = "B2"
    else:
        tier = "C"
    return {
        "role_family": fam,
        "must_keywords": "、".join(must),
        "plus_keywords": "、".join(plus),
        "must_hit_keywords": "、".join(must_hit),
        "plus_hit_keywords": "、".join(plus_hit),
        "must_hit_rate": round(must_rate, 4),
        "ats_score": round(ats_score, 2),
        "scene_score": round(scene_score, 2),
        "evidence_score": round(evidence_score, 2),
        "value_score": round(value_score, 2),
        "score": total,
        "tier": tier,
        "hit_points": hit_points,
        "gap_points": gap_points,
        "resume_version": family_to_resume_version(fam),
    }


def family_to_resume_version(family: str) -> str:
    if family == "数据开发/数据工程":
        return "数据开发_定制版"
    if family == "策略分析/增长分析":
        return "策略增长_定制版"
    return "数据分析商分_定制版"


def delivery_priority(tier: str, score: float, has_return_offer: bool, days_old: int) -> float:
    tier_base = {"A1": 100, "A2": 90, "B1": 70, "B2": 60, "C": 30}.get(tier, 20)
    return tier_base + score * 0.2 + (8 if has_return_offer else 0) + max(0, 20 - days_old) * 0.1


def classify_evidence_strength(must_hit_rate: float, quant_count: int, plus_hit_keywords: str) -> str:
    plus_hit_cnt = len([x for x in str(plus_hit_keywords).split("、") if x.strip()])
    if must_hit_rate >= 0.8 and (quant_count >= 6 or plus_hit_cnt >= 2):
        return "strong"
    if must_hit_rate >= 0.5 and (quant_count >= 3 or plus_hit_cnt >= 1):
        return "medium"
    return "weak"


def compute_competition_score(row: pd.Series, company_job_count: int) -> Dict[str, object]:
    tier_weight = {"A1": 35, "A2": 30, "B1": 22, "B2": 15, "C": 8}
    city = str(row.get("city", ""))
    city_weight = 25 if "上海" in city else 12
    company_weight = min(30, 8 + company_job_count * 1.8)
    score = round(tier_weight.get(str(row.get("tier", "C")), 8) + city_weight + company_weight, 2)
    level = "high" if score >= 70 else ("medium" if score >= 50 else "low")
    return {"competition_score": score, "competition_level": level}


def build_resume_documents(profile: Dict[str, object], top_jobs: pd.DataFrame, gap_df: pd.DataFrame) -> None:
    common_doc = Document()
    common_doc.add_heading("张靖恒-秋招大厂数据岗简历（通用优化版）", level=1)
    common_doc.add_paragraph(f"可到岗时间：尽快 | 每周可实习：{profile['intern_days']}天 | 可实习时长：{profile['intern_months']}个月 | 毕业年份：{profile['grad_year']}届")
    common_doc.add_paragraph("技能关键词优先级：" + "、".join(profile["skills"][:18]))
    common_doc.add_heading("岗位高频关键词补齐建议", level=2)
    for _, r in gap_df.head(12).iterrows():
        common_doc.add_paragraph(f"- 补齐关键词：{r['missing_keyword']}（影响岗位数：{r['affected_jobs']}）")
    common_doc.add_heading("项目描述改写模板", level=2)
    common_doc.add_paragraph("动作动词 + 技术方法 + 业务场景 + 量化结果 + 业务价值")
    common_doc.add_paragraph("示例：搭建SQL+Python自动化分析链路，服务电商增长场景，周报生成效率提升60%，支撑活动策略迭代。")
    common_doc.save(os.path.join(DATA_DIR, "张靖恒-秋招大厂数据岗简历_通用优化版.docx"))

    folder = os.path.join(DATA_DIR, "张靖恒-秋招大厂数据岗简历_分岗位族定制版文件夹")
    versions = {
        "数据分析商分_定制版.docx": "重点突出：指标体系、多维分析、可视化、业务闭环与增长结果。",
        "数据开发_定制版.docx": "重点突出：ETL、数据建模、数仓链路、调度稳定性与数据质量。",
        "策略增长_定制版.docx": "重点突出：策略实验、用户分层、留存转化、AB测试与收益归因。",
    }
    for fname, desc in versions.items():
        d = Document()
        d.add_heading(fname.replace(".docx", ""), level=1)
        d.add_paragraph(f"可到岗时间：尽快 | 每周可实习：{profile['intern_days']}天 | 可实习时长：{profile['intern_months']}个月")
        d.add_paragraph(desc)
        d.add_heading("优先投递岗位（Top 10）", level=2)
        sub = top_jobs[top_jobs["resume_version"] == fname.replace(".docx", "")]
        if sub.empty:
            sub = top_jobs.head(10)
        for _, r in sub.head(10).iterrows():
            d.add_paragraph(f"- {r['company']}｜{r['name']}｜评分{r['score']}｜命中：{r['hit_points']}")
        d.save(os.path.join(folder, fname))


def main() -> None:
    ensure_dirs()
    resume_text = read_resume_text(RESUME_PATH)
    profile = extract_resume_profile(resume_text)
    target = pd.read_csv(TARGET_PATH, keep_default_na=False)
    master = pd.read_csv(MASTER_PATH, keep_default_na=False)
    false_negative = pd.read_csv(FALSE_NEGATIVE_PATH, keep_default_na=False) if os.path.exists(FALSE_NEGATIVE_PATH) else pd.DataFrame()
    all_jobs = pd.concat([target, false_negative], ignore_index=True).drop_duplicates(subset=["company", "name", "url"], keep="first")
    all_jobs["hard_pass"], all_jobs["hard_reason"] = zip(*all_jobs.apply(lambda r: hard_screen_row(r, profile), axis=1))
    all_jobs["required_fields_missing"] = all_jobs.apply(
        lambda r: int(
            sum(
                1
                for c in ["name", "company", "city", "publish_time", "deadline", "jd_raw", "url"]
                if str(r.get(c, "")).strip() == ""
            )
        ),
        axis=1,
    )
    all_jobs["field_quality"] = all_jobs["required_fields_missing"].map(lambda x: "high" if x <= 1 else ("medium" if x <= 3 else "low"))
    all_jobs.to_csv(os.path.join(REPORT_DIR, "internship_post_screening_result.csv"), index=False, encoding="utf-8-sig")

    passed = all_jobs[all_jobs["hard_pass"]].copy()
    score_parts = passed.apply(lambda r: match_score(r, profile), axis=1, result_type="expand")
    dup_cols = [c for c in score_parts.columns if c in passed.columns]
    if dup_cols:
        passed = passed.drop(columns=dup_cols)
    passed = pd.concat([passed.reset_index(drop=True), score_parts.reset_index(drop=True)], axis=1)
    passed["publish_dt"] = passed["publish_time"].map(normalize_publish_date)
    passed["days_old"] = passed["publish_dt"].map(lambda d: max(0, (datetime.now() - d).days))
    passed["has_return_offer"] = passed["jd_raw"].map(lambda x: any(k in str(x) for k in RETURN_OFFER_KEYWORDS))
    passed["delivery_priority"] = passed.apply(
        lambda r: delivery_priority(r["tier"], float(r["score"]), bool(r["has_return_offer"]), int(r["days_old"])),
        axis=1,
    )
    passed = passed.sort_values(["tier", "delivery_priority", "score"], ascending=[True, False, False])
    topn = passed.copy()
    topn["投递建议"] = topn["tier"].map(
        {
            "A1": "立即投递（优先内推）",
            "A2": "直接投递（对应定制版）",
            "B1": "补1个关键词后投递",
            "B2": "微调项目描述后投递",
            "C": "暂不投递",
        }
    )
    a_shortlist = topn[topn["tier"].isin(["A1", "A2"])].copy().head(30)
    a_shortlist = a_shortlist.assign(a_batch_rank=range(1, len(a_shortlist) + 1))
    topn = topn.merge(a_shortlist[["company", "name", "url", "a_batch_rank"]], on=["company", "name", "url"], how="left")
    topn["is_a_batch_recommended"] = topn["a_batch_rank"].notna()
    topn["evidence_strength"] = topn.apply(
        lambda r: classify_evidence_strength(float(r["must_hit_rate"]), int(profile["quant_count"]), str(r["plus_hit_keywords"])),
        axis=1,
    )
    company_job_counts = topn.groupby("company").size().to_dict()
    comp_parts = topn.apply(
        lambda r: compute_competition_score(r, int(company_job_counts.get(str(r["company"]), 0))),
        axis=1,
        result_type="expand",
    )
    topn = pd.concat([topn.reset_index(drop=True), comp_parts.reset_index(drop=True)], axis=1)
    topn.to_csv(os.path.join(REPORT_DIR, "resume_job_match_topN.csv"), index=False, encoding="utf-8-sig")

    sentence_rows = []
    for _, r in topn.head(300).iterrows():
        jd = f"{r.get('responsibility','')}。{r.get('requirement','')}。{r.get('jd_raw','')}"
        for item in extract_sentence_mapping(jd):
            sentence_rows.append(
                {
                    "company": r["company"],
                    "job_name": r["name"],
                    "url": r["url"],
                    "sentence_role": item["role"],
                    "sentence": item["sentence"],
                    "skills": item["skills"],
                }
            )
    sentence_df = pd.DataFrame(sentence_rows)
    sentence_df.to_csv(os.path.join(REPORT_DIR, "jd_sentence_mapping.csv"), index=False, encoding="utf-8-sig")

    time_constraint_rows = []
    for _, r in all_jobs.iterrows():
        txt = f"{r.get('name','')} {r.get('jd_raw','')} {r.get('requirement','')}"
        extracted = extract_time_constraints(txt, default_grad_year=int(profile["grad_year"]))
        time_constraint_rows.append(
            {
                "company": r.get("company", ""),
                "job_name": r.get("name", ""),
                "url": r.get("url", ""),
                "intern_days_required": extracted["intern_days_required"],
                "intern_months_required": extracted["intern_months_required"],
                "experience_years_required": extracted["experience_years_required"],
                "grad_year_required": extracted["grad_year_required"],
                "arrival_required": extracted["arrival_required"],
                "is_hard_pass": bool(r.get("hard_pass", False)),
            }
        )
    pd.DataFrame(time_constraint_rows).to_csv(os.path.join(REPORT_DIR, "time_constraints_report.csv"), index=False, encoding="utf-8-sig")

    term_vocab = {}
    for s in topn["must_keywords"].fillna("").tolist() + topn["plus_keywords"].fillna("").tolist():
        for k in [x.strip() for x in str(s).split("、") if x.strip()]:
            term_vocab[k] = int(term_vocab.get(k, 0)) + 1
    pd.DataFrame(
        [{"term": k, "count": v} for k, v in sorted(term_vocab.items(), key=lambda x: x[1], reverse=True)]
    ).to_csv(os.path.join(REPORT_DIR, "normalized_term_vocab.csv"), index=False, encoding="utf-8-sig")

    summary = (
        topn.groupby(["company", "role_family"], as_index=False)
        .agg(
            jobs=("name", "count"),
            avg_score=("score", "mean"),
            a1=("tier", lambda s: int((s == "A1").sum())),
            a2=("tier", lambda s: int((s == "A2").sum())),
            b=("tier", lambda s: int((s.isin(["B1", "B2"])).sum())),
        )
        .sort_values(["a1", "a2", "avg_score"], ascending=[False, False, False])
    )
    summary["avg_score"] = summary["avg_score"].round(2)
    summary.to_csv(os.path.join(REPORT_DIR, "company_role_fit_summary.csv"), index=False, encoding="utf-8-sig")

    a_jobs = topn[topn["tier"].isin(["A1", "A2", "B1", "B2"])].copy()
    must_counter = Counter()
    for s in a_jobs["must_keywords"].fillna("").tolist():
        for k in [x for x in str(s).split("、") if x]:
            must_counter[k] += 1
    resume_skill = set(profile["skills"])
    gap_rows = []
    for k, c in must_counter.most_common(50):
        if k not in resume_skill:
            gap_rows.append({"missing_keyword": k, "affected_jobs": c, "suggestion": f"在相关项目经历补充 {k} 的实际应用场景与结果"})
    gap_df = pd.DataFrame(gap_rows)
    gap_df.to_csv(os.path.join(REPORT_DIR, "resume_gap_analysis.csv"), index=False, encoding="utf-8-sig")

    contribution_rows = []
    for _, r in topn.iterrows():
        must_hit = [x for x in str(r["must_hit_keywords"]).split("、") if x]
        plus_hit = [x for x in str(r["plus_hit_keywords"]).split("、") if x]
        denom = max(1, len(must_hit) * 1.4 + len(plus_hit) * 1.0)
        for s in must_hit:
            contribution_rows.append(
                {
                    "company": r["company"],
                    "job_name": r["name"],
                    "tier": r["tier"],
                    "skill": s,
                    "evidence_type": "must",
                    "contribution_rate": round(1.4 / denom, 4),
                    "evidence_strength": r["evidence_strength"],
                }
            )
        for s in plus_hit:
            contribution_rows.append(
                {
                    "company": r["company"],
                    "job_name": r["name"],
                    "tier": r["tier"],
                    "skill": s,
                    "evidence_type": "plus",
                    "contribution_rate": round(1.0 / denom, 4),
                    "evidence_strength": r["evidence_strength"],
                }
            )
    contribution_df = pd.DataFrame(contribution_rows)
    contribution_df.to_csv(os.path.join(REPORT_DIR, "skill_contribution_report.csv"), index=False, encoding="utf-8-sig")

    comp_high = topn[topn["competition_level"] == "high"].copy()
    comp_features = (
        comp_high.groupby(["company", "role_family"], as_index=False)
        .agg(jobs=("name", "count"), avg_score=("score", "mean"), avg_competition=("competition_score", "mean"))
        .sort_values(["jobs", "avg_competition"], ascending=[False, False])
    )
    comp_features["avg_score"] = comp_features["avg_score"].round(2)
    comp_features["avg_competition"] = comp_features["avg_competition"].round(2)
    comp_features["strategy"] = comp_features.apply(
        lambda x: "48小时内投递+补齐must证据" if int(x["jobs"]) >= 3 else "优先内推并强化量化成果",
        axis=1,
    )
    comp_features.to_csv(os.path.join(REPORT_DIR, "high_competition_strategy_report.csv"), index=False, encoding="utf-8-sig")

    a_for_calendar = a_jobs[a_jobs["tier"].isin(["A1", "A2", "B1"])].copy().head(60)
    plans = []
    current = TODAY
    count = 0
    for _, r in a_for_calendar.iterrows():
        if count and count % 6 == 0:
            current = current + timedelta(days=1)
        plans.append(
            {
                "plan_date": str(current),
                "company": r["company"],
                "job_name": r["name"],
                "tier": r["tier"],
                "priority": round(float(r["delivery_priority"]), 2),
                "resume_version": r["resume_version"],
                "micro_tuning": f"补齐关键词：{r['gap_points']}",
                "url": r["url"],
            }
        )
        count += 1
    pd.DataFrame(plans).to_csv(os.path.join(REPORT_DIR, "delivery_calendar.csv"), index=False, encoding="utf-8-sig")

    prep_rows = []
    for _, r in a_jobs[a_jobs["tier"].isin(["A1", "A2"])].head(80).iterrows():
        fam = r["role_family"]
        if fam == "数据开发/数据工程":
            ask = "数仓分层、ETL设计、数据质量、调度故障恢复"
        elif fam == "策略分析/增长分析":
            ask = "增长漏斗、策略实验、归因分析、ROI评估"
        elif fam == "数据科学/算法分析":
            ask = "特征工程、模型评估、离在线一致性、业务落地"
        else:
            ask = "指标体系、业务分析闭环、可视化表达、跨部门协作"
        prep_rows.append(
            {
                "company": r["company"],
                "job_name": r["name"],
                "tier": r["tier"],
                "core_focus": ask,
                "evidence_answer": f"优先讲述：{r['hit_points']} 对应项目成果",
                "gap_fix": f"补充表达：{r['gap_points']}",
            }
        )
    pd.DataFrame(prep_rows).to_csv(os.path.join(REPORT_DIR, "interview_prep_guide.csv"), index=False, encoding="utf-8-sig")

    md_lines = []
    md_lines.append("# 张靖恒-秋招简历岗位专属改写建议")
    md_lines.append("")
    md_lines.append("## A/B档岗位5分钟微调模板")
    for _, r in topn[topn["tier"].isin(["A1", "A2", "B1", "B2"])].head(50).iterrows():
        md_lines.append(f"### {r['company']}｜{r['name']}｜{r['tier']}｜评分{r['score']}")
        md_lines.append(f"- 命中点：{r['hit_points']}")
        md_lines.append(f"- 缺口点：{r['gap_points']}")
        md_lines.append(f"- 微调动作：将 {r['gap_points']} 写入最相关项目 bullet，并补1条量化结果")
        md_lines.append(f"- 推荐简历版本：{r['resume_version']}")
        md_lines.append(f"- 链接：{r['url']}")
        md_lines.append("")
    with open(os.path.join(DATA_DIR, "张靖恒-秋招简历_岗位专属改写建议.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    build_resume_documents(profile, topn, gap_df)

    profile_df = pd.DataFrame(
        [
            {
                "intern_days": profile["intern_days"],
                "intern_months": profile["intern_months"],
                "grad_year": profile["grad_year"],
                "skills_count": len(profile["skills"]),
                "scene_count": len(profile["scenes"]),
                "quant_count": profile["quant_count"],
            }
        ]
    )
    profile_df.to_csv(os.path.join(REPORT_DIR, "resume_profile_snapshot.csv"), index=False, encoding="utf-8-sig")

    metrics = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "jobs_total": int(len(all_jobs)),
        "jobs_passed": int(len(passed)),
        "topn_rows": int(len(topn)),
        "evidence_strength_counts": topn["evidence_strength"].value_counts().to_dict(),
        "competition_level_counts": topn["competition_level"].value_counts().to_dict(),
        "sentence_role_counts": sentence_df["sentence_role"].value_counts().to_dict() if not sentence_df.empty else {},
    }
    with open(os.path.join(REPORT_DIR, "jd_decision_metrics.json"), "w", encoding="utf-8") as f:
        json.dump(metrics, f, ensure_ascii=False, indent=2)

    print("done", len(all_jobs), len(passed), len(topn))


if __name__ == "__main__":
    main()
